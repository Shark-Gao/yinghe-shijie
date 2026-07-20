import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "docxdeps"))

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs", "硬核视界_多平台数据基线与复盘手册_2026-07-17.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "666666"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_borders(table, color="D9E1EA"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_paragraph_format(paragraph, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def add_text(paragraph, text, size=11, color="000000", bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_body(doc, text, before=0, after=6):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=before, after=after, line=1.25)
    add_text(p, text)
    return p


def add_labelled(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=4, line=1.25)
    add_text(p, label, size=10.5, color=INK, bold=True)
    add_text(p, text, size=10.5, color="222222")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_format(p, before=18, after=10, line=1.0, keep=True)
        add_text(p, text, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_paragraph_format(p, before=14, after=7, line=1.0, keep=True)
        add_text(p, text, size=13, color=BLUE, bold=True)
    else:
        set_paragraph_format(p, before=10, after=5, line=1.0, keep=True)
        add_text(p, text, size=12, color=DARK_BLUE, bold=True)
    return p


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_borders(table)
    header_cells = table.rows[0].cells
    for idx, value in enumerate(headers):
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
        p = header_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, before=0, after=0, line=1.1)
        add_text(p, value, size=9.2, color=INK, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            set_paragraph_format(p, before=0, after=0, line=1.18)
            if idx in (0, 1) and len(value) <= 15:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, size=font_size, color="222222")
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    set_borders(table, color="D5DEE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=2, after=2, line=1.2)
    add_text(p, title + "  ", size=10.5, color=INK, bold=True)
    add_text(p, body, size=10.5, color="222222")
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_format(p, before=0, after=0, line=1.0)
    add_text(p, "硬核视界｜内部复盘基线｜", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.0)
    add_text(p, "硬核视界  /  多平台数据基线与复盘手册", size=8.5, color=MUTED)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header(section)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    # Page 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p, before=12, after=4, line=1.0)
    add_text(p, "硬核视界", size=11, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=4, line=1.0)
    add_text(p, "多平台数据基线与复盘手册", size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=18, line=1.15)
    add_text(p, "首期基线：截至 2026 年 7 月 17 日  |  用于后续数据对比与剪辑策略复盘", size=11, color=MUTED)

    add_callout(doc, "这份文档怎么用", "每次复盘先补齐同一统计周期内的原始数据，再对照本次基线看变化。不要只比较播放量；先看初始分发、观看留存、互动价值和涨粉转化四段链路。")

    add_heading(doc, "本期总判断", 1)
    add_body(doc, "账号已经证明硬核知识题材能够获得初始分发：小红书有高收藏和高涨粉，抖音在同类作者中播放、互动和涨粉均高于平均，快手获得了最大规模的测试流量。", after=5)
    add_body(doc, "三端共同短板是观看留存。当前多个作品为 10 分钟左右的完整解说，但平均观看仅 15 至 47 秒；长版内容与短视频消费场景不匹配，限制了平台继续扩大推荐。", after=8)

    add_heading(doc, "平台角色与当前优先级", 2)
    add_table(doc,
              ["平台", "当前最强信号", "当前短板", "建议角色"],
              [
                  ["小红书", "收藏与涨粉质量最高", "完播率低", "沉淀关注与可收藏知识"],
                  ["抖音", "同类对比中播放、互动、涨粉领先；有搜索流量", "完播低于同类", "测试爆款钩子与选题"],
                  ["快手", "播放与平台助推最强", "完播、互动、转粉最弱", "短版本分发实验"],
              ],
              [1250, 3000, 2300, 2810])

    add_heading(doc, "数据使用边界", 2)
    add_labelled(doc, "统计周期：", "截图中同时存在近 7 日、近 30 日和单条作品数据；跨平台仅比较结构和趋势，不把绝对播放量当成统一排名。")
    add_labelled(doc, "账号阶段：", "样本仍处于起号早期，单条作品和单日峰值会显著影响整体结果。后续至少以连续 10 至 15 条短版本为一个测试周期。")
    add_labelled(doc, "口径提醒：", "小红书“观看数”、抖音/快手“播放量”及不同后台的互动指标不完全等价；更新时请保留后台原始名称。")

    doc.add_page_break()

    # Page 2
    add_heading(doc, "首期基线数据", 1)
    add_body(doc, "来源：用户提供的各平台后台截图。大部分截图统计截止至 2026 年 7 月 17 日；下表保留原后台口径。", after=8)
    add_table(doc,
              ["平台", "统计范围", "播放/观看", "留存", "互动", "涨粉/转粉", "分发或补充信号"],
              [
                  ["小红书", "近 30 日", "曝光 2.4 万；观看 5,317", "平均 39 秒；完播 1.9%", "赞 205；评 6；藏 159；分享 18", "涨粉 70；主页转粉 8.5%", "封面点击率 8.5%；视频推荐 65%，首页推荐 30%"],
                  ["抖音", "近 7 日 / 近 30 日", "近 7 日 5,345；近 30 日 6,289", "近 7 日完播 2.7%；近 30 日 2.26%", "互动率 3.1%（同类 2.6%）；赞 139", "近 7 日净涨粉 29（同类 10）", "搜索量 831；近 7 日播放高于同类 64.44%"],
                  ["快手", "近 7 日 / 近 30 日", "近 7 日 3.1 万；近 30 日 3.9 万", "完播 1.4%", "赞 54；评 14；分享 0", "作者净涨粉 18；作品涨粉 2", "近 7 日播放超过 90.7% 同类；额外助推 6,243 次"],
              ],
              [760, 1100, 1460, 1280, 1650, 1450, 1660],
              font_size=8.3)

    add_heading(doc, "可复用的解读知识", 2)
    add_labelled(doc, "点击/封面：", "小红书封面点击率 8.5%，说明题材和包装已经能让用户进入内容；抖音个别单条点击率也不错。下一步应优先优化进入后的前 10 秒，而不是先大改封面。")
    add_labelled(doc, "留存：", "完播率需要结合时长判断。10 分钟作品平均观看几十秒，核心问题通常不是知识密度不够，而是答案出现太晚、背景铺垫太长、一个视频承担了太多问题。")
    add_labelled(doc, "互动价值：", "收藏、分享和关注通常比单纯点赞更能反映内容的长期价值。小红书的收藏和涨粉说明“可保存的硬核知识”已形成正反馈；快手的低分享则提示内容尚未形成明确转述价值。")
    add_labelled(doc, "分发判断：", "有初始推荐但曲线迅速回落，常见原因是前段留存和互动不足。平台给流量不是结论，是否形成二次扩推才是验证。")

    doc.add_page_break()

    # Page 3
    add_heading(doc, "平台复盘结论", 1)
    add_heading(doc, "小红书：最值得优先沉淀的高质量受众", 2)
    add_body(doc, "观看 5,317 带来 70 位新增关注，且有 159 次收藏。对于起号期，这说明观众不仅看了，而且认为内容值得留下或继续追踪。当前首要任务是提升完播，而不是改变题材。", after=5)
    add_labelled(doc, "适配表达：", "45 至 90 秒、单一问题、强信息密度、可收藏结论。标题优先突出结果、规模、成本、稀缺性或现实用途。")
    add_labelled(doc, "保留要素：", "封面图形化、题目明确、工程可视化、系列化选题。二战战机等军工题材已显示出强转粉潜力。")

    add_heading(doc, "抖音：适合测试钩子、搜索词和大众化选题", 2)
    add_body(doc, "近 7 日播放、互动率和净涨粉均高于同类作者，但完播率低于同类。说明选题能进推荐池，叙事节奏还没有接住用户。", after=5)
    add_labelled(doc, "已有线索：", "“一条运河如何改写全球航线”比抽象芯片讲解更易获得播放；其优势是问题大、结果明确、普通人无需前置知识。")
    add_labelled(doc, "适配表达：", "30 至 60 秒，第一秒直接给冲突或反常识结果；保留 SR-71、巴拿马运河、芯片等明确搜索词。")

    add_heading(doc, "快手：曝光测试有效，内容转化尚未建立", 2)
    add_body(doc, "快手取得了最高播放和明显的平台助推，但 3.9 万播放对应的赞、分享、作品涨粉都较少，且完播率最低。当前不能把高播放直接理解为账号已跑通。", after=5)
    add_labelled(doc, "适配表达：", "更直给的口语化问题、更大的字幕、更早出现的结论和画面变化。先用 30 至 60 秒版本测试，不宜继续直接搬运 10 分钟长版。")
    add_labelled(doc, "判断方式：", "下一周期同时看完播、点赞、评论、分享和作品涨粉是否一起上升；只有播放上涨而转化不动，仍属于低质量分发。")

    doc.add_page_break()

    # Page 4
    add_heading(doc, "下一轮优化方案", 1)
    add_callout(doc, "核心策略", "一源多剪、平台分别包装。完整解说保留给更适合长内容的平台；小红书、抖音、快手只用同一源视频中最有冲突、最有画面的一个问题做短版本。")

    add_heading(doc, "短视频结构", 2)
    add_table(doc,
              ["时间段", "内容任务", "写法示例"],
              [
                  ["0 至 2 秒", "给结果、冲突或异常画面", "“导弹为什么追不上 SR-71？”"],
                  ["2 至 10 秒", "说明这个结果为什么重要", "“它不是单纯飞得快，而是能让拦截系统失去机会。”"],
                  ["10 至 45 秒", "只解释一个关键机制", "用一张结构图或一个动态过程讲清原因。"],
                  ["结尾", "给结论，并埋下一条内容的问题", "“下一条讲它为什么飞到机身会发烫。”"],
              ],
              [1200, 2600, 5560],
              font_size=9.3)

    add_heading(doc, "选题与剪辑原则", 2)
    add_labelled(doc, "优先选题：", "军工装备、超级工程、产业竞争、全球航线、AI 与芯片。优先“一个大众问题 + 一个震撼答案”的组合。")
    add_labelled(doc, "避免方式：", "一个视频同时讲完整历史、所有原理和所有细节；先花几十秒铺背景、最后才说答案；各平台用同一条 10 分钟成片直接分发。")
    add_labelled(doc, "拆条方法：", "将一条 10 分钟源视频拆为 5 至 8 条。可按速度、成本、内部结构、关键部件、失败风险、现实影响分别切题，每条只保留一个答案。")
    add_labelled(doc, "互动设计：", "结尾不泛泛要求点赞，而提出可回答的问题，例如“你觉得这种技术最难的是材料、发动机还是控制系统？”并在置顶评论补充一个未讲完的关键事实。")

    add_heading(doc, "本轮测试目标", 2)
    add_table(doc,
              ["平台", "连续测试内容", "优先提升指标", "保留指标"],
              [
                  ["小红书", "10 条 45 至 90 秒单问题版本", "完播率、平均观看时长", "收藏率、播放转粉"],
                  ["抖音", "10 条 30 至 60 秒强钩子版本", "完播率、前 10 秒留存", "搜索量、互动率、净涨粉"],
                  ["快手", "10 条 30 至 60 秒口语化版本", "完播率、分享、作品涨粉", "平台助推后的互动转化"],
              ],
              [1000, 3100, 2600, 2660])

    doc.add_page_break()

    # Page 5
    add_heading(doc, "后续复盘填写页", 1)
    add_body(doc, "每次复盘请尽量选同一统计周期（建议近 7 日或近 30 日）并记录到数据截止日。若发布数量明显不同，需同时记录作品数。", after=8)

    add_heading(doc, "本期信息", 2)
    add_table(doc,
              ["字段", "填写内容"],
              [
                  ["复盘周期", "____________________________"],
                  ["数据截止日", "____________________________"],
                  ["本期测试主题", "____________________________"],
                  ["作品数量", "____________________________"],
                  ["版本说明", "例如：45 秒短版 / 60 秒强钩子版 / 长版拆条"],
              ],
              [2600, 6760])

    add_heading(doc, "平台数据对比", 2)
    add_table(doc,
              ["平台", "播放/观看", "完播率", "平均观看", "互动", "收藏/分享", "涨粉", "与基线相比"],
              [
                  ["小红书", "", "", "", "", "", "", ""],
                  ["抖音", "", "", "", "", "", "", ""],
                  ["快手", "", "", "", "", "", "", ""],
              ],
              [900, 1200, 1050, 1150, 1200, 1250, 1050, 1560])

    add_heading(doc, "复盘结论", 2)
    add_labelled(doc, "本期最佳选题：", "____________________________________________________________")
    add_labelled(doc, "有效钩子或开头：", "________________________________________________________")
    add_labelled(doc, "主要流失节点：", "__________________________________________________________")
    add_labelled(doc, "下期只保留的做法：", "______________________________________________________")
    add_labelled(doc, "下期要停止的做法：", "______________________________________________________")
    add_labelled(doc, "下期要验证的假设：", "______________________________________________________")

    add_callout(doc, "复盘原则", "一次只验证一到两个变量：例如只改开头，或只改时长。若同时改选题、封面、时长和剪辑节奏，就很难判断真正带来变化的因素。")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.core_properties.title = "硬核视界多平台数据基线与复盘手册"
    doc.core_properties.subject = "短视频多平台数据基线与持续复盘"
    doc.core_properties.author = "硬核视界"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
